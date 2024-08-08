from docx import Document
from django.http import HttpResponse

def generate_docx(proposal_text):
    # Create a new Document
    doc = Document()
    
    # Add a title
    doc.add_heading('Website Development Proposal', level=1)
    
    # Remove asterisks from the proposal text
    proposal_text = proposal_text.replace('*', '')  # Remove asterisks

    # Split the proposal into paragraphs based on double new lines
    paragraphs = proposal_text.split('\n\n')  # Adjust this based on your proposal structure

    for paragraph in paragraphs:
        # Check if it's a title
        if paragraph.startswith("Title:"):  # Adjust based on your title format
            title = paragraph.replace("Title:", "").strip()
            doc.add_heading(title, level=2)  # Use level 2 for subheadings
        else:
            # Add each regular paragraph
            doc.add_paragraph(paragraph.strip())

    # Save the document to an in-memory file
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="proposal.docx"'
    doc.save(response)

    return response

